import json
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

def set_heading(doc, text, level):
    heading = doc.add_heading(text, level=level)
    return heading

def parse_and_add_json_to_doc(doc, data, level=1):
    if isinstance(data, dict):
        for key, value in data.items():
            formatted_key = key.replace('_', ' ').title()
            if isinstance(value, (dict, list)):
                doc.add_heading(formatted_key, level=level)
                parse_and_add_json_to_doc(doc, value, level + 1)
            else:
                p = doc.add_paragraph()
                p.add_run(f"{formatted_key}: ").bold = True
                p.add_run(str(value))
    elif isinstance(data, list):
        for item in data:
            if isinstance(item, (dict, list)):
                parse_and_add_json_to_doc(doc, item, level)
            else:
                doc.add_paragraph(str(item), style='List Bullet')
    else:
        doc.add_paragraph(str(data))

def generate_startup_document(project, agent_results) -> str:
    """Generates a DOCX file and returns the file path."""
    doc = Document()
    
    # Title
    title = doc.add_heading("Startup Pitch & Strategy Document", 0)
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    
    # Intro
    doc.add_heading("Project Overview", level=1)
    p = doc.add_paragraph()
    p.add_run("Idea: ").bold = True
    p.add_run(project.idea_description)
    
    p = doc.add_paragraph()
    p.add_run("Industry: ").bold = True
    p.add_run(project.industry)
    
    p = doc.add_paragraph()
    p.add_run("Target Audience: ").bold = True
    p.add_run(project.target_audience)
    
    doc.add_page_break()
    
    AGENT_TITLES = {
        'market_research': 'Market Research',
        'competitor_analysis': 'Competitor Analysis',
        'positioning': 'Brand Positioning',
        'landing_page': 'Landing Page Copy',
        'ad_copy': 'Ad Copy Generation',
        'email_marketing': 'Email Marketing Strategy'
    }
    
    for result in agent_results:
        agent_type = result.agent_type
        title_text = AGENT_TITLES.get(agent_type, agent_type.replace('_', ' ').title())
        
        doc.add_heading(title_text, level=1)
        
        output_data = result.output_data
        parse_and_add_json_to_doc(doc, output_data, level=2)
        doc.add_page_break()
        
    import tempfile
    import os
    temp_dir = tempfile.gettempdir()
    file_path = os.path.join(temp_dir, f"startup_{project.id}.docx")
    doc.save(file_path)
    return file_path
